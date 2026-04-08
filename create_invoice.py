import csv
import sys
import os
from datetime import datetime
from docx import Document


if len(sys.argv) < 2:
    print("Missing client ID")
    sys.exit(1)

client_id = sys.argv[1]

# paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
csv_file = os.path.join(BASE_DIR, "data", "clients.csv")
template_path = os.path.join(BASE_DIR, "assets", "template.docx")
output_dir = os.path.join(BASE_DIR, "invoices")

counter_file = os.path.join(BASE_DIR, "invoice_counter.txt")

os.makedirs(output_dir, exist_ok=True)


client_name = None
already_invoiced = False
transactions = []
all_rows = []


# ---- READ CSV ----
with open(csv_file, newline='') as file:
    reader = csv.reader(file)
    for row in reader:
        if len(row) < 6:
            all_rows.append(row)
            continue

        id_, name, date, work, amount, invoiced = row

        if id_ == client_id:
            client_name = name
            if invoiced.upper() == 'NO':
                transactions.append((date, work, float(amount)))
                row[5] = "YES"  # mark as invoiced
            else:
                pass

        all_rows.append(row)

if not transactions:
    print("No transactions found.")
    sys.exit(1)

#if already_invoiced:
    #print(f"Client has been invoiced for all balances.")
   # sys.exit(0)

with open(csv_file, "w", newline='') as file:
    writer = csv.writer(file)
    writer.writerows(all_rows)

#DEFINE INVOICE_NUMBER BEFORE PRINT
if not os.path.exists(counter_file):
    with open(counter_file, "w") as f:
        f.write("1000")
    current_invoice_num = 1000
else:
    with open(counter_file, "r") as f:
        current_invoice_num = int(f.read().strip()) + 1

# Save the new number back to the file
with open(counter_file, "w") as f:
    f.write(str(current_invoice_num))

invoice_number = str(current_invoice_num)

print(f"Invoice {invoice_number} generated for {client_name}!")


# ---- LOAD TEMPLATE ----
doc = Document(template_path)

# Replacing place holders
invoice_date = datetime.now().strftime("%m/%d/%Y")

for paragraph in doc.paragraphs:
    if "[CLIENT NAME]" in paragraph.text:
        paragraph.text = paragraph.text.replace("[CLIENT NAME]", client_name)

    if "[INVOICE DATE]" in paragraph.text:
        paragraph.text = paragraph.text.replace("[INVOICE DATE]", invoice_date)

    if "[INVOICE NUMBER]" in paragraph.text:
        paragraph.text = paragraph.text.replace("[INVOICE NUMBER]", invoice_number)

# Looking at table
table = doc.tables[0]  # assumes first table is your service table

# Remove existing empty rows except header
while len(table.rows) > 1:
    table._element.remove(table.rows[1]._element)

# Adding transactions
total = 0

for date, work, amount in transactions:
    row_cells = table.add_row().cells
    row_cells[0].text = work
    row_cells[1].text = date
    row_cells[2].text = f"${amount:.2f}"
    total += amount

# Total row
row_cells = table.add_row().cells
row_cells[0].text = "TOTAL"
row_cells[1].text = ""
row_cells[2].text = f"${total:.2f}"

# Save file path
filename = f"invoice_{invoice_number}.docx"
filepath = os.path.join(output_dir, filename)

doc.save(filepath)



